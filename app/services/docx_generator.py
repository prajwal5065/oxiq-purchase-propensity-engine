"""DOCX Report Generator - compiles a company dossier analysis into a
professional Word document format (.docx) directly from backend state.
"""
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from app.models.company import Company
from app.schemas.explanation import AnalysisExplanation


def set_cell_background(cell, fill_hex: str):
    """Set the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner margins (padding) for a table cell (in twentieths of a point, dxas)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def add_heading_styled(doc, text: str, level: int):
    """Add a heading with custom size, spacing, and color."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level > 1 else 18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.bold = True
    
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(26, 36, 56)  # Dark Navy
        # Add a subtle bottom border or line below H1
        p_border = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '46505F')  # Slate Gray
        p_border.append(bottom)
        p._p.get_or_add_pPr().append(p_border)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(70, 80, 95)  # Slate Gray
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(100, 100, 100)
    
    return p


def generate_company_docx(company: Company, explanation: AnalysisExplanation) -> io.BytesIO:
    """Generate a beautifully styled report document from the company dossier data."""
    doc = Document()
    
    # 1. Page Setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Configure default style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(51, 51, 51)  # Off-black
    
    # 2. Header & Title Block
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    run_title = title_p.add_run("OxiQ Purchase Propensity Report")
    run_title.bold = True
    run_title.font.size = Pt(22)
    run_title.font.color.rgb = RGBColor(26, 36, 56)  # Dark Navy
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(18)
    run_subtitle = subtitle_p.add_run(f"Dossier Analysis for {company.name}")
    run_subtitle.italic = True
    run_subtitle.font.size = Pt(12)
    run_subtitle.font.color.rgb = RGBColor(120, 120, 120)
    
    # 3. Metadata Table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = 'Light Shading Accent 1'
    meta_data = [
        ("Company Name", company.name),
        ("Website Domain", company.domain),
        ("Industry Sector", company.industry or "Not Classified"),
        ("Analysis Timestamp", company.last_processed_at.strftime("%Y-%m-%d %H:%M:%S UTC") if company.last_processed_at else "Not processed"),
    ]
    for i, (label, val) in enumerate(meta_data):
        row = meta_table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = val
        set_cell_background(row.cells[0], "F2F4F7")
        set_cell_margins(row.cells[0], top=80, bottom=80, left=120, right=120)
        set_cell_margins(row.cells[1], top=80, bottom=80, left=120, right=120)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # 4. Propensity Verdict Block
    add_heading_styled(doc, "Purchase Propensity Verdict", level=1)
    
    # Get purchase score
    purchase_score = 0.0
    purchase_scores = [s for s in company.scores if s.score_type == "purchase_propensity"]
    if purchase_scores:
        purchase_score = max(purchase_scores, key=lambda s: s.created_at).value

    decision = explanation.disqualification.final_decision  # qualified / disqualified / insufficient_data
    category_label = explanation.disqualification.category.value.replace("_", " ").title()
    
    verdict_table = doc.add_table(rows=1, cols=3)
    verdict_table.autofit = False
    
    # Score Cell
    c_score = verdict_table.rows[0].cells[0]
    c_score.width = Inches(2.2)
    p_score_lbl = c_score.paragraphs[0]
    p_score_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_score_lbl = p_score_lbl.add_run("SCORE\n")
    run_score_lbl.font.size = Pt(10)
    run_score_lbl.font.color.rgb = RGBColor(100, 100, 100)
    run_score_val = p_score_lbl.add_run(f"{purchase_score:.0f}")
    run_score_val.bold = True
    run_score_val.font.size = Pt(36)
    run_score_val.font.color.rgb = RGBColor(26, 36, 56)
    run_score_tot = p_score_lbl.add_run("/100")
    run_score_tot.font.size = Pt(12)
    run_score_tot.font.color.rgb = RGBColor(120, 120, 120)
    set_cell_background(c_score, "F8F9FA")
    set_cell_margins(c_score, top=140, bottom=140, left=100, right=100)
    
    # Verdict Cell
    c_verdict = verdict_table.rows[0].cells[1]
    c_verdict.width = Inches(2.2)
    p_verd_lbl = c_verdict.paragraphs[0]
    p_verd_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_verd_lbl = p_verd_lbl.add_run("VERDICT\n")
    run_verd_lbl.font.size = Pt(10)
    run_verd_lbl.font.color.rgb = RGBColor(100, 100, 100)
    
    run_verd_val = p_verd_lbl.add_run(decision.upper().replace("_", " "))
    run_verd_val.bold = True
    run_verd_val.font.size = Pt(18)
    
    # Color coding based on decision
    if decision == "qualified":
        run_verd_val.font.color.rgb = RGBColor(34, 139, 34)  # Forest Green
        bg_color = "E6F4EA"
    elif decision == "disqualified":
        run_verd_val.font.color.rgb = RGBColor(178, 34, 34)  # Firebrick Red
        bg_color = "FCE8E6"
    else:
        run_verd_val.font.color.rgb = RGBColor(120, 120, 120)  # Slate Gray
        bg_color = "F1F3F4"
        
    set_cell_background(c_verdict, bg_color)
    set_cell_margins(c_verdict, top=140, bottom=140, left=100, right=100)

    # Coverage/Confidence Cell
    c_metrics = verdict_table.rows[0].cells[2]
    c_metrics.width = Inches(2.1)
    p_met = c_metrics.paragraphs[0]
    p_met.paragraph_format.line_spacing = 1.25
    
    cov_pct = explanation.evidence_coverage.coverage_percentage * 100
    conf_pct = explanation.confidence_explanation.overall_confidence * 100
    conf_level = explanation.confidence_explanation.level.title()
    
    run_cov = p_met.add_run("Evidence Coverage: ")
    run_cov.font.size = Pt(10)
    run_cov_val = p_met.add_run(f"{cov_pct:.0f}%\n")
    run_cov_val.bold = True
    run_cov_val.font.size = Pt(10)
    
    run_conf = p_met.add_run("Confidence Level: ")
    run_conf.font.size = Pt(10)
    run_conf_val = p_met.add_run(f"{conf_pct:.0f}% ({conf_level})\n")
    run_conf_val.bold = True
    run_conf_val.font.size = Pt(10)
    
    run_cat = p_met.add_run("Category: ")
    run_cat.font.size = Pt(10)
    run_cat_val = p_met.add_run(category_label)
    run_cat_val.bold = True
    run_cat_val.font.size = Pt(10)
    
    set_cell_background(c_metrics, "F8F9FA")
    set_cell_margins(c_metrics, top=140, bottom=140, left=120, right=120)
    
    # Verdict summary text
    p_summary = doc.add_paragraph()
    p_summary.paragraph_format.space_before = Pt(8)
    p_summary.paragraph_format.space_after = Pt(12)
    run_p_sum = p_summary.add_run(f"Headline: {explanation.headline}\n")
    run_p_sum.bold = True
    p_summary.add_run(f"Reasoning: {explanation.disqualification.primary_reason}")
    
    # 5. Executive Recommendations
    add_heading_styled(doc, "Executive Summary & Outreach Strategy", level=1)
    
    latest_recs = company.recommendations
    if latest_recs:
        rec = max(latest_recs, key=lambda r: r.created_at)
        
        # Summary Box
        p_rec_sum_title = doc.add_paragraph()
        p_rec_sum_title.paragraph_format.space_before = Pt(6)
        p_rec_sum_title.paragraph_format.space_after = Pt(2)
        run_sum_title = p_rec_sum_title.add_run("Executive Opportunity Summary")
        run_sum_title.bold = True
        run_sum_title.font.size = Pt(12)
        
        p_rec_sum = doc.add_paragraph(rec.executive_summary)
        p_rec_sum.paragraph_format.space_after = Pt(12)
        p_rec_sum.paragraph_format.left_indent = Inches(0.2)
        
        # Approach Box
        p_rec_app_title = doc.add_paragraph()
        p_rec_app_title.paragraph_format.space_before = Pt(6)
        p_rec_app_title.paragraph_format.space_after = Pt(2)
        run_app_title = p_rec_app_title.add_run("Recommended Outreach Approach")
        run_app_title.bold = True
        run_app_title.font.size = Pt(12)
        
        p_rec_app = doc.add_paragraph(rec.suggested_approach)
        p_rec_app.paragraph_format.space_after = Pt(12)
        p_rec_app.paragraph_format.left_indent = Inches(0.2)
        
        # Fit Reasons Bullets
        if rec.fit_reasons:
            p_fit_title = doc.add_paragraph()
            p_fit_title.paragraph_format.space_before = Pt(6)
            p_fit_title.paragraph_format.space_after = Pt(2)
            run_fit_title = p_fit_title.add_run("Key Fit Reasons")
            run_fit_title.bold = True
            
            for reason in rec.fit_reasons:
                p_bullet = doc.add_paragraph(style='List Bullet')
                p_bullet.paragraph_format.space_after = Pt(2)
                p_bullet.add_run(reason)
                
        # Top Risks Bullets
        if rec.top_risks:
            p_risk_title = doc.add_paragraph()
            p_risk_title.paragraph_format.space_before = Pt(6)
            p_risk_title.paragraph_format.space_after = Pt(2)
            run_risk_title = p_risk_title.add_run("Top Buying Gaps & Risks")
            run_risk_title.bold = True
            
            for risk in rec.top_risks:
                p_bullet = doc.add_paragraph(style='List Bullet')
                p_bullet.paragraph_format.space_after = Pt(2)
                p_bullet.add_run(risk)
    else:
        doc.add_paragraph("No recommendations generated yet for this company.")
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # 6. Pillar Scoring Breakdown
    add_heading_styled(doc, "Pillar Attribution Breakdown", level=1)
    
    pillar_table = doc.add_table(rows=1, cols=4)
    pillar_table.style = 'Medium Shading 1 Accent 1'
    
    headers = ["Pillar", "Score", "Confidence", "Findings / Rationale"]
    hdr_widths = [Inches(1.5), Inches(0.8), Inches(1.0), Inches(3.2)]
    
    # Format Headers
    hdr_cells = pillar_table.rows[0].cells
    for k, text in enumerate(headers):
        hdr_cells[k].text = text
        hdr_cells[k].paragraphs[0].runs[0].bold = True
        hdr_cells[k].width = hdr_widths[k]
        set_cell_background(hdr_cells[k], "1A2438")  # Dark Navy
        set_cell_margins(hdr_cells[k], top=100, bottom=100, left=100, right=100)
        
    # Get all individual scores
    for score in company.scores:
        if score.score_type == "purchase_propensity":
            continue
        row_cells = pillar_table.add_row().cells
        
        # Title
        p_name = score.score_type.value.replace("_", " ").title()
        row_cells[0].text = p_name
        row_cells[0].paragraphs[0].runs[0].bold = True
        
        # Score
        row_cells[1].text = f"{score.value:.0f}/100"
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Confidence
        conf_val = score.confidence or 0.0
        row_cells[2].text = f"{conf_val * 100:.0f}%"
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Findings
        reasons_text = ", ".join(score.reasons) if isinstance(score.reasons, list) else str(score.reasons or "")
        row_cells[3].text = reasons_text
        
        # Set widths and padding
        for idx in range(4):
            row_cells[idx].width = hdr_widths[idx]
            set_cell_margins(row_cells[idx], top=80, bottom=80, left=100, right=100)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # 7. Evidence Citations List
    add_heading_styled(doc, "Citations of Structured Evidence", level=1)
    
    if company.evidence_items:
        evidence_table = doc.add_table(rows=1, cols=4)
        evidence_table.style = 'Medium Shading 1 Accent 1'
        
        ev_headers = ["Pillar", "Source / Category", "Date / Conf.", "Evidence Details / Excerpt"]
        ev_widths = [Inches(1.2), Inches(1.5), Inches(1.0), Inches(2.8)]
        
        # Format Headers
        hdr_cells = evidence_table.rows[0].cells
        for k, text in enumerate(ev_headers):
            hdr_cells[k].text = text
            hdr_cells[k].paragraphs[0].runs[0].bold = True
            hdr_cells[k].width = ev_widths[k]
            set_cell_background(hdr_cells[k], "1A2438")  # Dark Navy
            set_cell_margins(hdr_cells[k], top=100, bottom=100, left=100, right=100)
            
        for ev in sorted(company.evidence_items, key=lambda x: (x.pillar.value if x.pillar else "", x.source.value)):
            row_cells = evidence_table.add_row().cells
            
            # Pillar
            row_cells[0].text = ev.pillar.value.replace("_", " ").title() if ev.pillar else "Unassigned"
            row_cells[0].paragraphs[0].runs[0].bold = True
            
            # Source
            src_str = ev.source.value.replace("_", " ").title() if ev.source else "Unknown"
            cat_str = f" ({ev.category})" if ev.category else ""
            row_cells[1].text = f"{src_str}{cat_str}"
            
            # Date / Conf
            date_str = ev.published_at.strftime("%Y-%m-%d") if ev.published_at else "No Date"
            conf_val = ev.confidence or 1.0
            row_cells[2].text = f"{date_str}\n[Conf: {conf_val * 100:.0f}%]"
            row_cells[2].paragraphs[0].runs[0].font.size = Pt(9.5)
            row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(100, 100, 100)
            
            # Details
            p_details = row_cells[3].paragraphs[0]
            p_details.paragraph_format.line_spacing = 1.15
            
            # Excerpt/Text description
            payload_text = ""
            if isinstance(ev.payload, dict):
                # Try to summarize key values from structured payload
                p_items = []
                for key, val in ev.payload.items():
                    if key in ("matched_terms", "technologies", "technology"):
                        p_items.append(f"{key}: {val}")
                    elif key == "excerpt" and val:
                        payload_text = f'"{val}"'
                    elif isinstance(val, (str, int, float)) and key not in ("markdown", "html"):
                        p_items.append(f"{key}: {val}")
                if p_items and not payload_text:
                    payload_text = "; ".join(p_items)
            else:
                payload_text = str(ev.payload or "")
                
            if payload_text:
                run_pay = p_details.add_run(payload_text + "\n\n")
                run_pay.italic = True
                run_pay.font.size = Pt(10)
                
            if ev.url:
                run_url = p_details.add_run(f"Source URL: {ev.url}")
                run_url.font.size = Pt(9)
                run_url.font.color.rgb = RGBColor(70, 80, 95)
                
            # Set widths and padding
            for idx in range(4):
                row_cells[idx].width = ev_widths[idx]
                set_cell_margins(row_cells[idx], top=85, bottom=85, left=100, right=100)
    else:
        doc.add_paragraph("No evidence items cited for this company.")
        
    # Write to IO stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
